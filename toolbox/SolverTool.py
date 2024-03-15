from langchain.chains import LLMChain
from langchain.chat_models import ChatOpenAI
from docx import Document
import configparser
from opencc import OpenCC
import base64
from PIL import Image
import pytesseract
from toolbox.Loader import *
from toolbox.gpt_image import GPTimage

config = configparser.ConfigParser()
config.read('config.ini')
VISION_MODEL = config.get('openai', 'vision_model')
OUTPUT_FOLDER = config.get('folder', 'output_folder')

class SolverTool():
    def __init__(self, key: list, solver_prompt, cc) -> None:
        self.cc = cc
        self.solver_agent = GPTimage(
            openai_api_key = key[0],
            prompt = solver_prompt,
            model_name = VISION_MODEL
        )  

    def encode_image(self, image_path):
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')  
        
    def img2text(self, path, language):
        img = Image.open(path)
        pytesseract.pytesseract.tesseract_cmd = config.get('pytesseract', 'pytesseract')
        if language == 'en':
            lang = 'eng'
        elif language == 'ch':
            lang = 'chi_tra'
        else: 
            return -1
        text = pytesseract.image_to_string(img, lang=lang) # eng
        # print(text)
        return text
        
    def run(self, img_path: str, language: str) -> str:
        text = self.img2text(img_path, language)
        if text == -1:
            return "The language is not supported."     
        if language == 'ch':
            result = self.cc.convert(self.solver_agent.run({"language": "中文","discription": text},  [self.encode_image(img_path)]))
        else: 
            result = self.solver_agent.run({"language": "English","discription": text},  [self.encode_image(img_path)])
        return result

def solve_problem(language = 'ch'):
    if config.get('pytesseract', 'pytesseract').split('.')[1].find('exe') == -1:
        print("pytesseract not found")
        return -3
    agent = SolverTool(
        key = [config.get('openai', 'key1')],
        solver_prompt=load_prompt("solver_prompt.txt"),
        cc=OpenCC('s2twp'),
    )
    inputList = os.listdir(INPUT_FOLDER)
    if (inputList[0].find('.jpg') == -1 and inputList[0].find('.png') == -1) or len(inputList) == 0:
        return -1
    path = INPUT_FOLDER + '/' + inputList[0]
    print("Solving problem............")
    result = agent.run(path, language)
    if result == 'gpt error':
        return -2
    print("Writing answers............")
    doc = Document()
    doc.add_picture(path)
    par = doc.add_paragraph(result)
    # for para in doc.paragraphs:
    #     if para.style.name.startswith('Heading'):
    #         continue       
    #     result = agent.run(para.text)
    #     para.text = result
    doc.save(OUTPUT_FOLDER + '/' + inputList[0][:-4] + "_answer.docx")
    return 0