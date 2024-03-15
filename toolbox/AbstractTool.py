from langchain.chains import LLMChain
from langchain.chat_models import ChatOpenAI
from docx import Document
import configparser
from opencc import OpenCC
from toolbox.Loader import *

config = configparser.ConfigParser()
config.read('config.ini')
GPT_MODEL = config.get('openai', 'model')
OUTPUT_FOLDER = config.get('folder', 'output_folder')

class AbstractTool():
    def __init__(self, key: list, abstract_prompt, cc) -> None:
        self.cc = cc
        self.abstract_agent = LLMChain(
            llm=ChatOpenAI(temperature=0, model_name=GPT_MODEL, openai_api_key=key[0]),
            prompt=abstract_prompt,
            verbose=False
        )

    def run(self, language: str, word: int, text: str):
        result = self.cc.convert(self.abstract_agent.run({"language": language,"word": word,"content": text}))
        return result

def abstract(language='en', word=200):
    if language == 'ch':
        lang = '中文'
    else:
        lang = 'English'
    agent = AbstractTool(
        key = [config.get('openai', 'key1')],
        abstract_prompt=load_prompt("abstract_prompt.txt"),
        cc=OpenCC('s2twp')
    )
    inputList = os.listdir(INPUT_FOLDER)
    pdfList = []
    wordList = []
    for input in inputList:
        if input.find('.pdf') != -1:
            pdfList.append(input)
        elif input.find('.docx') != -1:
            wordList.append(input)
    texts = load_pdf()
    print("Condensing abstract............")
    i = 0
    for text in pdfList:
        result = agent.run(lang, word, text)
        doc = Document()
        doc.add_heading("Abstract", 0)
        doc.add_paragraph(result)
        doc.save(OUTPUT_FOLDER + '/' + pdfList[i].split('.')[0] + "_abstract.docx")
        percentage = 100*(i+1)/len(inputList)
        msg = f"File {i+1} abstract completed {percentage}%."
        print(msg)
        i=i+1
    texts = load_word()
    for text in texts:
        result = agent.run(lang, text, word)
        doc = Document()
        doc.add_heading("Abstract", 0)
        doc.add_paragraph(result)
        doc.save(OUTPUT_FOLDER + '/' + wordList[i].split('.')[0] + "_abstract.docx")
        percentage = 100*(i+1+len(pdfList))/len(inputList)
        msg = f"File {i+1} abstract completed {percentage}%."
        print(msg)
        i=i+1
    return 0