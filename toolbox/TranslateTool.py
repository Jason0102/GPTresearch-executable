from langchain.chains import LLMChain
from langchain.chat_models import ChatOpenAI
from docx import Document
import configparser
from opencc import OpenCC
from toolbox.Loader import *

config = configparser.ConfigParser()
config.read('config.ini')
GPT_MODEL = config.get('openai', 'model')
OUTPUT_FOLDER = config.get('folder', 'output_folder')

class TranslateTool():
    def __init__(self, key: list, translate_prompt, cc) -> None:
        self.cc = cc
        self.translate_agent = LLMChain(
            llm=ChatOpenAI(temperature=0, model_name=GPT_MODEL, openai_api_key=key[0]),
            prompt=translate_prompt,
            verbose=False
        )
    
    def run(self, domain, language, text) -> str:
        if language == 'ch':
            lang = "中文"
        elif language == 'en':
            lang = "English"
        elif language == 'de':
            lang = "Deutsch"
        elif language == 'jp':
            lang = "日本語"
        else:
            lang = "中文"

        if language == 'ch':
            result = self.cc.convert(self.translate_agent.run({"domain": domain, "language": lang, "content": text}))
        else:
            result = self.translate_agent.run({"domain": domain, "language": lang, "content": text})
        return result

def professinal_translation(domain="自動控制", language='ch'):
    # text = "Arckmann's formula is an approach for feedback controller design, which can arbitrary assigned the pole if the system is controllable."
    if language != 'ch' and language != 'en' and language != 'jp' and language != 'de':
        result =  "This language is not supported."   
        doc = Document()
        doc.add_heading("Translated content",1)
        par = doc.add_paragraph(result)
        doc.save(OUTPUT_FOLDER + '/' + inputList[i][:-4] + '_' + language + ".docx")
        return -1
    agent = TranslateTool(
        key = [config.get('openai', 'key1')],
        translate_prompt=load_prompt("translate_prompt.txt"),
        cc=OpenCC('s2twp')
    )
    inputList = os.listdir(INPUT_FOLDER)
    pdfList = []
    wordList = []
    for input in inputList:
        if input.find('.pdf') != -1:
            pdfList.append(input)
        elif input.find('.docx') != -1:
            wordList.append(input)
    print("Translating............")
    word_texts = load_word()
    if len(word_texts) != 0:
        i = 0
        for text in word_texts:
            paragraph = text.split('\n')
            doc = Document()
            doc.add_heading("Translated content",1)
            n = 0
            for p in paragraph:
                if p == "":
                    continue       
                result = agent.run(domain, language, p)
                par = doc.add_paragraph(result)
                percentage = 100*(n+1)/len(paragraph)
                msg = f"File {i+1} translating completed {percentage}%."
                print(msg)
                n = n+1
            nameroot = wordList[i][:-4]
            doc.save(OUTPUT_FOLDER + '/' + nameroot + '_' + language + ".docx")
            i = i+1
    return 0