from langchain.chains import LLMChain
from langchain.chat_models import ChatOpenAI
from docx import Document
import configparser
from opencc import OpenCC
from toolbox.Loader import *

config = configparser.ConfigParser()
config.read('config.ini')
GPT_MODEL = config.get('openai', 'model')
OUTPUT_FOLDER = config.get('folder', 'output_folder')

class PaperAnalyzeTool():
    def __init__(self, key: list, reader_prompt, ch_prompt, cc) -> None:
        self.cc = cc
        self.reader_agent = LLMChain(
            llm=ChatOpenAI(temperature=0, model_name=GPT_MODEL, openai_api_key=key[0]),
            prompt=reader_prompt,
            verbose=False
        )
        self.ch_agent = LLMChain(
            llm=ChatOpenAI(temperature=0, model_name=GPT_MODEL, openai_api_key=key[1]),
            prompt=ch_prompt,
            verbose=False
        )

    def run(self, text:str, chinese=True):
        result = self.reader_agent.run({"content": text}) # produce report
        if chinese:
            result_ch = self.cc.convert(self.ch_agent.run({"content": result})) # English to traditional chinese
        else:
            result_ch = None
        reply = [result, result_ch]
        # + "\n########################\n" + result_ch
        return reply

def to_docx(data: list, chinese: bool):
    filename="Reading"
    print("Writing report...............................")
    doc = Document()
    doc.add_heading(filename,0)
    i = 0
    for d in data[0]:
        i = i+1
        index = f"Paper {i}"
        doc.add_heading(index,1)
        doc.add_heading('Title: ',2)
        par = doc.add_paragraph(d["title"])
        # par.add_run(d["title"])
        doc.add_heading('Authors: ',2)
        par1 = doc.add_paragraph(d["authors"])
        doc.add_heading('Journal: ',2)
        par2 = doc.add_paragraph(d["journal"])
        doc.add_heading('Info: ',2)
        par3 = doc.add_paragraph(d["info"])
        doc.add_heading('Methods: ',2)
        par4 = doc.add_paragraph(d["methods"])
        doc.add_heading('Experiment: ',2)
        par5 = doc.add_paragraph(d["experiment"])
        doc.add_heading('Contribution: ',2)
        par6 = doc.add_paragraph(d["contribution"])
        doc.add_heading('Defect: ',2)
        par7 = doc.add_paragraph(d["defect"])
        # doc.add_page_break()
    doc.save(OUTPUT_FOLDER + '/' + filename + ".docx")
    if chinese:
        doc_ch = Document()
        doc_ch.add_heading(filename,0)
        i = 0
        for d in data[1]:
            i = i+1
            index = f"文章 {i}"
            doc_ch.add_heading(index,1)
            doc_ch.add_heading('標題：',2)
            par = doc_ch.add_paragraph(d["title"])
            # par.add_run(d["title"])
            doc_ch.add_heading('作者： ',2)
            par1 = doc_ch.add_paragraph(d["authors"])
            doc_ch.add_heading('期刊：',2)
            par2 = doc_ch.add_paragraph(d["journal"])
            doc_ch.add_heading('資訊：',2)
            par3 = doc_ch.add_paragraph(d["info"])
            doc_ch.add_heading('方法：',2)
            par4 = doc_ch.add_paragraph(d["methods"])
            doc_ch.add_heading('實驗：',2)
            par5 = doc_ch.add_paragraph(d["experiment"])
            doc_ch.add_heading('貢獻：',2)
            par6 = doc_ch.add_paragraph(d["contribution"])
            doc_ch.add_heading('缺陷： ',2)
            par7 = doc_ch.add_paragraph(d["defect"])
            # doc.add_page_break()
        doc_ch.save(OUTPUT_FOLDER + '/' + filename + "_ch.docx")
    return 0

def batch_analyze_paper(chinese=True):
    agent = PaperAnalyzeTool(
        key = [config.get('openai', 'key1'), config.get('openai', 'key2')],
        reader_prompt=load_prompt("reader_prompt.txt"),
        ch_prompt=load_prompt("ch_prompt.txt"),
        cc=OpenCC('s2twp')
    )
    papers = []
    papers = load_pdf() + load_word()
    if len(papers) == 0:
        return -1
    print("Amount of papers: ", len(papers))
    print("Reading paper.........................")
    data = [[],[]]
    i = 0
    for paper in papers:
        output = agent.run(paper)
        i = i+1
        # 英文版
        result = output[0]
        title = result[result.find('Title:')+7:result.find('Authors:')-2 ]
        authors = result[result.find('Authors:')+9:result.find('Journal:')-2 ]
        journal = result[result.find('Journal:')+9:result.find('Info:')-2] 
        info = result[result.find('Info:')+6:result.find('Methods:')-2 ]
        methods = result[result.find('Methods:')+9:result.find('Experiment:')-2] 
        experiment = result[result.find('Experiment:')+12:result.find('Contribution:')-2 ]
        contribution = result[result.find('Contribution:')+14:result.find('Defect:')-2]
        defect = result[result.find('Defect:')+8:]
        data[0].append({"title": title, "authors": authors, "journal": journal, "info": info, "methods": methods, "experiment": experiment, "contribution": contribution, "defect": defect})
        if chinese:# 中文版
            result = output[1]
            title = result[result.find('標題：')+3:result.find('作者：')-2 ]
            authors = result[result.find('作者：')+3:result.find('期刊：')-2 ]
            journal = result[result.find('期刊：')+3:result.find('資訊：')-2] 
            info = result[result.find('資訊：')+3:result.find('方法：')-2 ]
            methods = result[result.find('方法：')+3:result.find('實驗：')-2] 
            experiment = result[result.find('實驗：')+3:result.find('貢獻：')-2 ]
            contribution = result[result.find('貢獻：')+3:result.find('缺陷：')-2]
            defect = result[result.find('缺陷：')+3:]
            data[1].append({"title": title, "authors": authors, "journal": journal, "info": info, "methods": methods, "experiment": experiment, "contribution": contribution, "defect": defect})
        p = i/len(papers)*100
        print(f"paper analyzed {p} % ...............................")

    to_docx(data, chinese=chinese)
    return 0