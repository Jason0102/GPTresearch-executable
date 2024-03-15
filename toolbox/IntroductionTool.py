from langchain.chains import LLMChain
from langchain.chat_models import ChatOpenAI
from docx import Document
import configparser
from toolbox.Loader import *

config = configparser.ConfigParser()
config.read('config.ini')
GPT_MODEL = config.get('openai', 'model')
OUTPUT_FOLDER = config.get('folder', 'output_folder')

class IntroductionTool():
    def __init__(self, key: list, introduction_prompt, reference_prompt) -> None:
        self.introduction_agent = LLMChain(
            llm=ChatOpenAI(temperature=0, model_name='gpt-4-1106-preview', openai_api_key=key[0]),
            prompt=introduction_prompt,
            verbose=False
        )
        self.reference_agent = LLMChain(
            llm=ChatOpenAI(temperature=0, model_name='gpt-4-1106-preview', openai_api_key=key[0]),
            prompt=reference_prompt,
            verbose=False
        )

    def run(self, word, reference_type: str, text: str, mode='Sort alphabetically by last name.'):
        reference_list = self.reference_agent.run({"order": mode, "reference": text, "type": reference_type})
        result = self.introduction_agent.run({"word": word,"reference": text, "reference list": reference_list})
        return reference_list, result
    
    def rereference(self, reference_type: str, text: str, mode = 'Sort alphabetically by last name.'):
        reference_list = self.reference_agent.run({"order": mode, "reference": text, "type": reference_type})
        return reference_list
    
def reference_rearrange(reference_type='[1] J.H. Lo, "Spatiotemporal Features of Working Memory EEG," NTUME Journal, Vol. 1, No. 1, pp.1-13, 2024.'):
    config = configparser.ConfigParser()
    config.read('config.ini')
    agent = IntroductionTool(
        key=[config.get('openai', 'key1')],
        introduction_prompt=load_prompt("introduction_prompt.txt"),
        reference_prompt=load_prompt("reference_prompt.txt"),
    )
    fileList = os.listdir(INPUT_FOLDER)
    if fileList[0].find('.docx') == -1:
        return -1
    doc = Document(INPUT_FOLDER + '/' + fileList[0])
    text = ""
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            continue       
        text = text + para.text
    print("Generating reference list........................")
    re_list = agent.rereference(reference_type, text)
    doc = Document()
    doc.add_heading("Rereferences", 0)
    doc.add_paragraph(re_list)
    doc.save(OUTPUT_FOLDER + '/' + fileList[0].split('.')[0] + "_rereference.docx")

def write_introduction_by_reference(word=1000, reference_type='[1] J.H. Lo, "Spatiotemporal Features of Working Memory EEG," NTUME Journal, Vol. 1, No. 1, pp.1-13, 2024.'):
    config = configparser.ConfigParser()
    config.read('config.ini')
    agent = IntroductionTool(
        key=[config.get('openai', 'key1')],
        introduction_prompt=load_prompt("introduction_prompt.txt"),
        reference_prompt=load_prompt("reference_prompt.txt"),
    )
    fileList = os.listdir(INPUT_FOLDER)
    if fileList[0].find('.docx') == -1:
        return -1
    doc = Document(INPUT_FOLDER + '/' + fileList[0])
    text = ""
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            continue       
        text = text + para.text
    # reference_type = '[1] J.H. Lo, "Spatiotemporal Features of Working Memory EEG," NTUME Journal, Vol. 1, No. 1, pp.1-13, 2024.'
    print("Writing introduction with references......................................")
    re_list, result = agent.run(word, reference_type, text)
    doc = Document()
    doc.add_heading("Introduction with references", 0)
    doc.add_heading("Introduction", 1)
    doc.add_paragraph(result)
    doc.add_heading("reference", 1)
    doc.add_paragraph(re_list)
    doc.save(OUTPUT_FOLDER + '/' + fileList[0].split('.')[0] + "_introduction.docx")
    return 0