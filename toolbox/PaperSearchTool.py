from langchain.chains import LLMChain
from langchain.chat_models import ChatOpenAI
from docx import Document
import configparser
from scholarly import scholarly
from toolbox.Loader import *

config = configparser.ConfigParser()
config.read('config.ini')
GPT_MODEL = config.get('openai', 'model')
OUTPUT_FOLDER = config.get('folder', 'output_folder')

class PaperSearchTool():
    def __init__(self, key: list, prompt) -> None:
        self.keyword_agent = LLMChain(
            llm=ChatOpenAI(temperature=0, model_name=GPT_MODEL, openai_api_key=key[0]),
            prompt=prompt,
            verbose=False
        )

    def run(self, text: str, paper_N, keyword_N) -> str:
        try:
            keyword = self.keyword_agent.run({"human_input": text, "keyword_N": keyword_N}).split(',')
            output_list = []
            for j in range(keyword_N):
                search_query = scholarly.search_pubs(keyword[j])
                papers = [{"keyword": keyword[j]}]
                for i in range(paper_N):
                    info = next(search_query)
                    title = info['bib']['title']
                    author = info['bib']['author']
                    url = info['pub_url']
                    papers.append({'title': title, 'author': author, 'url': url})
                output_list.append(papers)
            return output_list
        except StopIteration:
            return []

def search_paper_by_paragraph(paper_N=1, keyword_N=1):
    # text = "Working memory is a complex cerebral function considered indicative of overall cognitive performance, regarded as the index of early stage of dementia. A card-pairing task is design to perform visual working memory with 59 older people participated. The research aims to discover the EEG features of older people during and provides the insight of cognitive aging.The working memory model is represented by concept of brain energy, connectivity, and complexity. Band power ratios describe the EEG waveforms in 5 frequencies. Brain connectivity is identified through magnitude squared coherence, phase locking value, and brain complexity is calculated by Katz fractal dimension. The results show that the anti-correlation of alpha and gamma wave forms can be observed in occipital lobe. Three pathways are revealed—the attention pathway, the short-term memory pathway, and the distraction resistance pathway. The changes of Katz fractal dimension are discovered in frontal and occipital lobe. Furthermore, working memory performance and cognitive ability can be predicted by the found spatiotemporal features. The research provides the insights of working memory related to other cognitive function, and the proposed working memory model can assess cognitive aging and realize the diagnosis of early stage of dementia."
    
    agent = PaperSearchTool(
        key = [config.get('openai', 'key1')],
        prompt = load_prompt("keyword_prompt_paper.txt")
    )
    inputList = os.listdir(INPUT_FOLDER)
    if (inputList[0].find('.pdf') == -1 and inputList[0].find('.docx') == -1) or len(inputList) == 0:
        return -1
    texts = []
    texts.append(load_pdf())
    texts.append(load_word())
    if len(texts) == 0:
        return -1
    print("Reading paper.........................")
    result = agent.run(texts[0], paper_N, keyword_N)
    
    print("Writing report...............................")
    doc = Document()
    doc.add_heading("References",0)
    i = 0
    for j in range(keyword_N):
        doc.add_heading('Keyword: ',2)
        par = doc.add_paragraph(result[j][0]['keyword'])
        for i in range(paper_N):
            doc.add_heading('Title: ',2)
            par = doc.add_paragraph(result[j][i+1]['title'])
            # par.add_run(d["title"])
            doc.add_heading('Authors: ',2)
            par1 = doc.add_paragraph(result[j][i+1]['author'])
            doc.add_heading('URL: ',2)
            par2 = doc.add_paragraph(result[j][i+1]['url'])
    output_file = inputList[0][:-4]
    doc.save(OUTPUT_FOLDER + '/' + output_file + "_reference.docx")